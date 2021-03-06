import os
import csv
import docx
import time
import requests
import pandas as pd
from pathlib import Path
from bs4 import BeautifulSoup
from setup import case, base
from general import pickle_file, create_project_dir
from general import dir_case, dir_Crawler, dir_scrapelist, dir_crawled
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize

directory_list = []
scrape_list = []
folder = case + '_scraper'
html = []
scrapedata = {}
stopwords_en = set(stopwords.words("english"))
stopwords_nl = set(stopwords.words("dutch"))


# =============================================================================
# Main scraper function
# =============================================================================


def scrape_case():
    create_tagslist()
    create_scrapelist()
    scrape_scrapelist()


# =============================================================================
# Supporting scraper functions
# =============================================================================


def create_tagslist():
    # go back to base directory to access html _text_elements.csv
    dir_Crawler(base)
    with open('html_text_elements.csv', newline='') as f:
        for row in csv.reader(f):
            html.append(row)


def create_scrapelist():
    scrape_links = []
    global scrape_list
    # go back to case crawled directory to store scrape results in crawl folder
    dir_crawled(base, case)

    # list the directories for each project folder
    for root, dirs, files in os.walk(os.getcwd(), topdown=False):
        for name in dirs:
            directory_list.append(os.path.join(root, name))

    # collect urls from all crawled.txt files and append to scrapelist
    for root, name in enumerate(directory_list):
        if (case + '_crawler') in name:
            file = Path.cwd().joinpath(name).joinpath("crawled.txt")
            with open(file) as f:
                for url in f:
                    scrape_links.append(url)

    # clear directory list, remove duplicates and go back to case folder
    directory_list.clear()
    scrape_list = list(dict.fromkeys(scrape_links))
    dir_case(base, case)

    # create scraper folder to save scrapelist to, then go back to temp dir
    create_project_dir(folder)
    dir_scrapelist(base, case)
    pickle_file("scrape_list.txt", scrape_list)
    dir_case(base, case)


# Code adopted from: 
# https://stackoverflow.com/questions/8733233/filtering-out-certain-bytes-in-python
def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (0x20 <= codepoint <= 0xD7FF or
            codepoint in (0x9, 0xA, 0xD) or
            0xE000 <= codepoint <= 0xFFFD or
            0x10000 <= codepoint <= 0x10FFFF)


def store_data(data, path):
    dir_case(base, case)
    doc = docx.Document()
    doc.add_heading(case, level=0)
    for url, pars in data.items():
        cleaned_string = ''.join(c for c in pars if valid_xml_char_ordinal(c))
        doc.add_heading(url, level=1)
        doc.add_paragraph(cleaned_string)
    doc.save(path)


def scrape_scrapelist():
    global scrape_list
    global scrapedata

    # attempt to scrape every url in scrape_list
    for index, url in enumerate(scrape_list):

        # create a soup for each url in scrape_list
        try:
            r = requests.get(url, timeout=1)
            r.raise_for_status()
            soup = BeautifulSoup(r.content, "html.parser")

            # retrieve text for all 'p' tags in the soup
            pars = [''.join(s.findAll(text=True)) for s in soup.findAll('p')]
            if not pars:
                for tag in enumerate(html):
                    for i in range(len(soup.find_all(tag))):
                        tags = [''.join(soup.find_all(tag)[i].get_text())]
                        scrapedata[url] = ' '.join(tags)
            else:
                scrapedata[url] = ' '.join(pars)
            time.sleep(10)

        except requests.exceptions.RequestException as err:
            print("RequestException: ", err)
            continue
        except requests.exceptions.HTTPError as errh:
            print("Http Error: ", errh)
            continue
        except requests.exceptions.ConnectionError as errc:
            print("Error Connecting: ", errc)
            continue
        except requests.exceptions.Timeout as errt:
            print("Timeout Error: ", errt)
            continue

    # tokenize sentences
    for url, text in scrapedata.items():
        scrapedata[url] = ' '.join(sent_tokenize(text))

    # store casedata to file in case directory
    store_data(scrapedata, "casedata.docx")
